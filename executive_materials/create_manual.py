from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "Manual_de_Uso_Open_Order_Control_Atualizado.docx"

PAPER = "FFFFFF"
GRAPHITE = "18181B"
INK = "52525B"
MIST = "F4F4F5"
ACCENT = "DC2626"
HAIRLINE = "E4E4E7"
WARNING = "FEF3C7"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(INK)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(GRAPHITE)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, after in [
        ("Title", 31, GRAPHITE, 16),
        ("Heading 1", 22, GRAPHITE, 14),
        ("Heading 2", 15, GRAPHITE, 8),
        ("Heading 3", 12, ACCENT, 6),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("OPEN ORDER CONTROL  |  MANUAL DE USO")
    set_run_font(run, size=8, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(1)
    line = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), ACCENT)
    borders.append(bottom)
    line.append(borders)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    add_page_number(p)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("—")
    set_run_font(r, size=22, color=ACCENT, bold=True)
    doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(13)
        r = p.add_run(subtitle)
        set_run_font(r, size=11.5, color=INK, italic=True)


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        set_run_font(r, bold=True, color=GRAPHITE)
        r = p.add_run(text[len(bold_start):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_note(doc, title, text, color=MIST):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, color)
    set_cell_border(cell, top={"val": "single", "sz": "6", "color": ACCENT}, bottom={"val": "single", "sz": "4", "color": HAIRLINE}, left={"val": "single", "sz": "4", "color": HAIRLINE}, right={"val": "single", "sz": "4", "color": HAIRLINE})
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper() + "  ")
    set_run_font(r, size=9, color=ACCENT, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=GRAPHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, title in enumerate(headers):
        cell = hdr.cells[idx]
        shade(cell, MIST)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top={"val": "single", "sz": "6", "color": ACCENT}, bottom={"val": "single", "sz": "4", "color": HAIRLINE}, left={"val": "single", "sz": "4", "color": HAIRLINE}, right={"val": "single", "sz": "4", "color": HAIRLINE})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        set_run_font(r, size=8.5, color=INK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_border(cell, top={"val": "single", "sz": "2", "color": HAIRLINE}, bottom={"val": "single", "sz": "2", "color": HAIRLINE}, left={"val": "single", "sz": "2", "color": HAIRLINE}, right={"val": "single", "sz": "2", "color": HAIRLINE})
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(value))
            set_run_font(r, size=9.1, color=GRAPHITE)
    if widths:
        set_column_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_steps(doc, steps):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    set_column_widths(table, [1.0, 14.5])
    for number, title, detail in steps:
        cells = table.add_row().cells
        shade(cells[0], GRAPHITE)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(number).zfill(2))
        set_run_font(r, size=11, color=PAPER, bold=True, name="Aptos Mono")
        shade(cells[1], MIST)
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title + " — ")
        set_run_font(r, size=10, color=GRAPHITE, bold=True)
        r = p.add_run(detail)
        set_run_font(r, size=10, color=INK)
        for cell in cells:
            set_cell_border(cell, top={"val": "single", "sz": "2", "color": PAPER}, bottom={"val": "single", "sz": "2", "color": PAPER}, left={"val": "single", "sz": "2", "color": PAPER}, right={"val": "single", "sz": "2", "color": PAPER})
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, image_path, caption, width=6.8):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(image_path), width=Inches(width))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(caption)
        set_run_font(r, size=8.5, color=INK, italic=True)
    else:
        add_note(doc, "Tela não localizada", f"A imagem esperada não foi encontrada: {image_path.name}", WARNING)


def create_crops():
    full = ASSETS / "open_order_control_dashboard_completo.png"
    if not full.exists():
        return []
    image = Image.open(full)
    width, height = image.size
    crops = []
    crop_specs = [
        ("manual_base_operacional.png", int(height * 0.38), min(height, int(height * 0.70))),
        ("manual_relatorios.png", int(height * 0.68), min(height, int(height * 0.96))),
    ]
    for file_name, top, bottom in crop_specs:
        if bottom - top < 100:
            continue
        crop = image.crop((0, top, width, bottom))
        output = ASSETS / file_name
        crop.save(output)
        crops.append(output)
    return crops


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_before = Pt(46)
    p = doc.add_paragraph()
    r = p.add_run("OPEN ORDER CONTROL")
    set_run_font(r, size=13, color=ACCENT, bold=True, name="Aptos Mono")
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    r = p.add_run("Manual de Uso")
    set_run_font(r, size=34, color=GRAPHITE, bold=True, name="Aptos Display")
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    r = p.add_run("Rotina operacional, métricas, controles e relatórios para gestão de pedidos em aberto")
    set_run_font(r, size=16, color=INK)
    p.paragraph_format.space_after = Pt(32)
    add_note(doc, "Finalidade", "Este documento orienta usuários operacionais, gestores e administradores na utilização completa da aplicação. As telas e recursos descritos refletem a versão validada em 15/08/2026.")
    doc.add_paragraph().paragraph_format.space_before = Pt(84)
    p = doc.add_paragraph()
    r = p.add_run("Público-alvo: Operação, Gestão, Diretoria e Administração")
    set_run_font(r, size=10.5, color=INK)
    p = doc.add_paragraph()
    r = p.add_run("Versão de referência: checkpoint fe76fa2a · Elaborado por Manus AI")
    set_run_font(r, size=9.5, color=INK)
    doc.add_page_break()


def build_manual():
    doc = Document()
    style_document(doc)
    create_crops()
    add_cover(doc)

    add_title(doc, "Como usar este manual", "Uma leitura orientada por objetivo: compreender, operar, analisar e reportar.")
    add_body(doc, "O Open Order Control foi concebido para transformar planilhas semanais de pedidos em aberto em uma base de acompanhamento histórico. A aplicação compara cada novo upload com os registros existentes, destaca mudanças de previsão de entrega, calcula sinais de risco e organiza a resposta por item, filial e exposição financeira.")
    add_table(doc, ["Capítulo", "Quando consultar", "Resultado esperado"], [
        ("1. Acesso e navegação", "Primeiro acesso ou mudança de perfil", "Entender o que é consulta e o que é função administrativa."),
        ("2. Importação semanal", "Quando receber a nova planilha", "Registrar um novo ciclo sem perder histórico."),
        ("3. Dashboard e riscos", "Reunião de acompanhamento", "Interpretar indicadores, alertas e prioridades."),
        ("4. Histórico e entregues", "Investigação de pedido", "Comprovar mudanças e status de carteira."),
        ("5. Relatórios", "Compartilhamento com áreas e diretoria", "Emitir PDF ou Excel no escopo adequado."),
    ], [3.2, 5.6, 7.0])
    add_note(doc, "Regra de ouro", "Nunca substitua dados manualmente entre semanas. A rastreabilidade depende da importação sequencial de cada planilha recebida.")
    doc.add_page_break()

    add_title(doc, "1. Visão geral e objetivos", "O que a aplicação controla e como ela apoia a gestão.")
    add_body(doc, "O sistema acompanha pedidos em aberto a partir de arquivos Excel importados semanalmente. A identificação do pedido é feita pela combinação Ship To + Item + Customer PO. Quando uma mesma combinação reaparece em uma nova carga, o sistema compara a previsão registrada no passado com a previsão atual e registra a mudança quando houver diferença.")
    add_table(doc, ["Objetivo", "Como o sistema responde"], [
        ("Preservar histórico", "Guarda a data da carga, arquivo, previsão anterior, previsão atual, variação em dias e quantidade acumulada de alterações."),
        ("Identificar risco", "Consolida instabilidade, vencimento, ausência de fornecedor, prioridade e exposição financeira em um Índice de Risco Executivo."),
        ("Organizar prioridade", "Classifica alertas, expõe fila de ação, sinais de decisão e pressão por filial."),
        ("Encerrar carteira corretamente", "Move para Itens Entregues os itens ativos ausentes no upload mais recente, sem excluir seu histórico."),
        ("Compartilhar informação", "Gera PDF Executivo, PDF Completo e Excel profissional com histórico completo ou apenas itens alterados."),
    ], [4.2, 11.6])
    doc.add_page_break()

    add_title(doc, "2. Perfis de acesso e controles", "Entenda quem pode consultar e quem pode alterar a base.")
    add_body(doc, "A aplicação distingue a consulta analítica das ações que modificam a base. O dashboard, filtros, históricos e relatórios podem ser analisados sem executar mudanças. As funções de upload e reset de importações exigem autenticação administrativa local.")
    add_table(doc, ["Recurso", "Disponibilidade", "Finalidade"], [
        ("Dashboard, filtros e históricos", "Consulta", "Acompanhar carteira, riscos e alterações."),
        ("Exportações PDF e Excel", "Consulta", "Compartilhar análises no escopo dos filtros ativos."),
        ("Upload semanal", "Administrador autenticado", "Registrar um novo ciclo e calcular comparações."),
        ("Resetar importações", "Administrador autenticado + confirmação", "Eliminar importações, itens e histórico em uma reinicialização controlada."),
        ("Configurações de Priorização", "Administrador autenticado", "Ajustar, simular, salvar ou restaurar os pesos usados no score da Fila de Ação."),
        ("Modo noturno", "Todos os usuários", "Ajustar a visualização; a preferência é mantida no navegador."),
    ], [4.4, 4.0, 7.4])
    add_note(doc, "Atenção", "O reset remove a base histórica de consulta. Execute-o somente quando houver decisão formal de reiniciar a operação e após salvar os relatórios necessários.", "FEF2F2")
    doc.add_page_break()

    add_title(doc, "3. Conheça a tela inicial", "Como ler o dashboard antes de partir para as exceções.")
    add_body(doc, "O dashboard reúne os principais dados para a tomada de decisão. Os indicadores superiores mostram o tamanho e a saúde da carteira; o Centro de Comando Estratégico combina índice de risco, sinais para decisão, pressão por filial e foco recomendado. A partir dele, o usuário pode aprofundar a análise nas abas e filtros inferiores.")
    add_figure(doc, ASSETS / "open_order_control_dashboard_topo.png", "Figura 1 — Visão superior do dashboard do Open Order Control. Fonte: aplicação validada no checkpoint 5f125413.")
    add_table(doc, ["Área da tela", "Como interpretar"], [
        ("KPIs superiores", "Mostram itens ativos, estabilidade, exposição financeira, vencidos, itens sem fornecedor e prioridades."),
        ("Centro de Comando", "Apresenta o Índice de Risco Executivo, recomendação de foco, sinais e pressão por filial."),
        ("Filtros", "Segmentam as visões por filial, item, PO, previsão e texto de busca."),
        ("Abas operacionais", "Permitem consultar Base Operacional, Mapa de Alterações e Itens Entregues."),
    ], [4.3, 11.5])
    doc.add_page_break()

    add_title(doc, "4. Preparar e importar a planilha semanal", "Roteiro para manter a comparação histórica confiável.")
    add_body(doc, "Antes de importar, confirme que a planilha contém a identificação do pedido, a filial ou endereço Ship To, o Customer PO, o item, a descrição quando disponível e a previsão de entrega. A consistência dessas colunas é essencial porque a comparação depende da chave Ship To + Item + Customer PO.")
    add_steps(doc, [
        (1, "Acesse como administrador", "Use o botão Entrar e autentique-se com o perfil ADM. As ações de upload e reset somente ficam disponíveis após esse acesso."),
        (2, "Revise o arquivo recebido", "Confirme se não há linhas de cabeçalho extras, células mescladas em meio aos dados ou alterações de nomenclatura nas colunas essenciais."),
        (3, "Selecione Upload de planilha semanal", "Escolha o arquivo Excel correspondente ao novo ciclo. Aguarde a conclusão do processamento antes de iniciar uma nova importação."),
        (4, "Leia o retorno da importação", "Verifique o total processado e, principalmente, se as alterações de previsão foram reconhecidas conforme esperado."),
        (5, "Analise os efeitos no dashboard", "Após a carga, consulte KPIs, alertas, pressão por filial, fila de ação e histórico dos itens relevantes."),
    ])
    add_note(doc, "Boa prática", "Mantenha o nome original do arquivo e uma rotina semanal definida. O nome do arquivo é preservado no histórico e ajuda na auditoria das alterações.")
    doc.add_page_break()

    add_title(doc, "5. Interpretar indicadores e o Índice de Risco Executivo", "A régua de priorização para a liderança.")
    add_body(doc, "O Índice de Risco Executivo varia de 0 a 100. Ele é uma leitura de priorização, não uma substituição da investigação do pedido. Sua composição ponderada privilegia a instabilidade e o vencimento, sem ignorar problemas de fornecedor, prioridades e exposição financeira.")
    add_note(doc, "Fórmula", "min(100, 35% × instabilidade + 25% × vencimento + 20% × sem fornecedor + 10% × prioridade + 10% × exposição financeira).")
    add_table(doc, ["Componente", "Peso", "Leitura gerencial"], [
        ("Instabilidade", "35%", "Frequência e concentração de alterações acumuladas de previsão."),
        ("Vencimento", "25%", "Pedidos cuja previsão está vencida e exigem ação imediata."),
        ("Sem fornecedor", "20%", "Risco de abastecimento ou ausência de informação de origem."),
        ("Prioridade", "10%", "Itens marcados como prioritários na carteira."),
        ("Exposição financeira", "10%", "Valor financeiro associado a itens sob risco."),
    ], [3.8, 2.0, 10.0])
    add_table(doc, ["Faixa", "Classificação", "Conduta recomendada"], [
        ("0–24", "Controlado", "Acompanhar a rotina e manter monitoramento semanal."),
        ("25–49", "Atenção", "Investigar as causas e verificar impacto por filial ou item."),
        ("50–100", "Crítico", "Tratar na fila de ação e levar a decisão gerencial."),
    ], [2.0, 3.1, 10.7])
    doc.add_page_break()

    add_title(doc, "6. Configurar e usar alertas de alteração", "Transforme mudança de prazo em exceção tratável.")
    add_body(doc, "O limiar de alerta é configurável em dias, entre 1 e 3.650, com valor inicial de 7 dias. Esse parâmetro define quais mudanças de previsão devem receber atenção visual. A severidade distingue um desvio relevante de um desvio que merece ação prioritária.")
    add_table(doc, ["Condição", "Classificação", "O que fazer"], [
        ("Variação menor ou igual ao limiar", "Sem alerta", "Monitorar normalmente na carteira."),
        ("Variação superior ao limiar", "Atenção", "Investigar motivo, fornecedor, impacto e nova data."),
        ("Variação igual ou superior ao dobro do limiar", "Crítico", "Priorizar em reunião, verificar plano de resposta e responsável."),
    ], [5.1, 3.4, 7.3])
    add_steps(doc, [
        (1, "Defina o limiar", "Ajuste o campo de dias conforme o nível de tolerância operacional da empresa."),
        (2, "Leia a proporção", "Use o gráfico de atenção versus crítico para entender o peso relativo das exceções."),
        (3, "Observe a tendência", "Acompanhe a evolução dos alertas críticos por upload para identificar deterioração ou recuperação."),
        (4, "Abra o histórico", "Para qualquer exceção importante, utilize Ver histórico antes de atribuir uma ação."),
    ])
    doc.add_page_break()

    add_title(doc, "7. Configurações de Priorização", "Como administrar a política de score da Fila de Ação.")
    add_body(doc, "A área Configurações de Priorização permite que o administrador ajuste os pontos atribuídos a cada critério que compõe o score operacional. A alteração é persistida e a Fila de Ação é recalculada imediatamente, sem necessidade de editar código ou reimportar arquivos. A funcionalidade foi criada para que a empresa alinhe a prioridade do sistema à sua política operacional vigente.")
    add_table(doc, ["Critério", "Peso padrão", "Efeito no score"], [
        ("Alteração de previsão", "4 pontos por alteração", "Eleva a prioridade proporcionalmente à instabilidade acumulada do item."),
        ("Sem fornecedor", "+5 pontos", "Destaca risco de abastecimento ou ausência de confirmação de origem."),
        ("Previsão vencida", "+3 pontos", "Evidencia itens cuja data planejada já foi superada."),
        ("Prioridade alta", "+2 pontos", "Acrescenta a relevância definida na origem do pedido."),
    ], [4.5, 3.2, 8.1])
    add_steps(doc, [
        (1, "Entre como administrador", "O botão Configurações aparece no cabeçalho após a autenticação administrativa."),
        (2, "Abra Configurações de Priorização", "Revise a explicação de cada critério e os valores atualmente ativos."),
        (3, "Ajuste os pesos", "Digite números inteiros não negativos de acordo com a política aprovada pela liderança."),
        (4, "Simule o resultado", "Observe a explicação de como a fila será recalculada antes de salvar a mudança."),
        (5, "Salve ou restaure", "Salvar persiste e aplica os pesos. Restaurar padrão retorna aos valores iniciais da aplicação."),
    ])
    add_note(doc, "Governança", "Altere os pesos somente após alinhamento com os responsáveis pela operação. A mudança afeta a ordenação de todos os itens na Fila de Ação e o balão explicativo do score passa a refletir os novos valores.")

    add_title(doc, "8. Filtrar a operação e analisar por filial", "Segmentação que se propaga para o dashboard e para os relatórios.")
    add_body(doc, "A aplicação usa o endereço Ship To para consolidar as filiais solicitantes. Para endereços específicos, regras de normalização convertem o endereço em cidades de referência, melhorando a consistência do filtro. Ao aplicar uma filial, o conjunto ativo do dashboard é segmentado; ao limpar, a visão retorna à carteira completa.")
    add_table(doc, ["Endereço reconhecido", "Filial consolidada"], [
        ("AVENIDA ASSIS BRASIL · RS", "Porto Alegre"),
        ("RUA ABEL SCUISSIATO · PR", "Colombo"),
        ("R VIDAL PROCOPIO LOHN · SC", "São José"),
        ("AV PREFEITO SINCLER SAMBATTI · PR", "Maringá"),
        ("RUA VALDEMIRO BELINSKI · SC", "Chapecó"),
    ], [10.0, 5.8])
    add_steps(doc, [
        (1, "Selecione a filial", "Use o filtro de filial solicitante para segmentar o conjunto de análise."),
        (2, "Aplique o filtro", "Confirme a aplicação para atualizar KPIs, tabelas, alertas e exportações."),
        (3, "Refine a busca", "Combine com item, PO, previsão ou busca textual para localizar uma situação específica."),
        (4, "Limpe quando terminar", "Use Limpar filtros para evitar conclusões baseadas em uma segmentação residual."),
    ])
    doc.add_page_break()

    add_title(doc, "9. Base Operacional, histórico e Mapa de Alterações", "Como investigar um item até o detalhe da mudança.")
    add_body(doc, "A Base Operacional reúne os itens ativos e oferece filtros de pesquisa. Além da previsão atual, ela apresenta a previsão anterior, a data do último upload, a última alteração, o total de mudanças e um resumo das alterações acumuladas. O Mapa de Alterações torna cada transição ainda mais explícita e o botão Ver histórico abre a linha do tempo do item.")
    add_figure(doc, ASSETS / "manual_base_operacional.png", "Figura 2 — Trecho da tela operacional, com áreas de consulta e acompanhamento. Fonte: aplicação Open Order Control.")
    add_table(doc, ["Campo", "Utilidade"], [
        ("Previsão atual", "Data atual de entrega, apresentada também em formato brasileiro nos relatórios Excel."),
        ("Previsão anterior", "Base imediata de comparação da última mudança."),
        ("Total de alterações", "Indica a recorrência de instabilidade do item."),
        ("Todas as datas de alteração", "Consolida as datas em que a previsão mudou."),
        ("Histórico das previsões", "Mostra a sequência de transições registrada para o pedido."),
    ], [5.2, 10.2])
    add_note(doc, "Roteiro de investigação", "Filtre o item ou PO → localize a linha → leia a previsão anterior e atual → abra Ver histórico → confirme o arquivo e a data de cada transição → registre a ação definida pela área responsável.")
    doc.add_page_break()

    add_title(doc, "10. Itens Entregues", "Como consultar a carteira encerrada sem perder rastreabilidade.")
    add_body(doc, "Para a aplicação, um item é considerado entregue quando estava ativo em um upload anterior e deixa de aparecer no upload mais novo. Essa regra foi definida para manter a carteira corrente limpa de registros encerrados, preservando as informações na aba Itens Entregues.")
    add_table(doc, ["Etapa", "Comportamento do sistema"], [
        ("Item consta no upload anterior", "Permanece ativo enquanto estiver presente no conjunto mais recente."),
        ("Item não consta no novo upload", "Recebe status de entregue e data de entrega conforme a identificação da ausência."),
        ("Consulta posterior", "Pode ser filtrado por filial, item e PO; os dados e histórico continuam disponíveis."),
        ("Leitura de KPIs", "O item deixa a carteira ativa, evitando distorção dos indicadores correntes."),
    ], [5.1, 10.3])
    add_note(doc, "Interpretação correta", "A classificação reflete a premissa operacional adotada: ausência em uma carga mais nova é tratada como entrega. Caso a regra de negócio da empresa mude, revise o processo antes de utilizar o indicador como confirmação física de recebimento.", WARNING)
    doc.add_page_break()

    add_title(doc, "11. Relatórios PDF e Excel", "Escolha o formato que corresponde ao objetivo da análise.")
    add_body(doc, "Todos os relatórios respeitam os filtros ativos no momento da exportação. Antes de emitir um arquivo, revise se a filial, item, PO, texto de busca ou demais recortes aplicados representam o público que receberá a informação.")
    add_figure(doc, ASSETS / "manual_relatorios.png", "Figura 3 — Trecho da área inferior da aplicação, com base operacional e comandos de exportação. Fonte: aplicação Open Order Control.")
    add_table(doc, ["Opção", "Conteúdo", "Quando usar"], [
        ("PDF Executivo", "Visão compacta de KPIs, risco, sinais, pressão e prioridades.", "Reuniões de liderança e compartilhamento rápido."),
        ("PDF Completo", "Todas as seções gerenciais, tendências, alertas, base e mapa.", "Análise abrangente e registro de reunião."),
        ("Exportar Excel", "Resumo Executivo, Base Operacional e Histórico de Alterações.", "Trabalho operacional, auditoria e análise em planilha."),
        ("Exportar alterados", "Apenas itens com alteração de previsão, com histórico completo.", "Follow-up de exceções e tratativa da semana."),
    ], [3.2, 6.4, 5.8])
    add_note(doc, "Excel profissional", "O arquivo contém cabeçalhos formatados, larguras ajustadas, filtros automáticos, painéis congelados, datas em formato brasileiro e uma aba dedicada ao Histórico de Alterações.")
    doc.add_page_break()

    add_title(doc, "12. Modo noturno e reset de importações", "Controles de uso e manutenção da base.")
    add_body(doc, "O modo noturno está disponível no menu superior. Ele é uma preferência de visualização e pode ser acionado ou desacionado sem alterar qualquer dado. A escolha é mantida no navegador, facilitando o uso em contextos de baixa luminosidade ou em longas jornadas de análise.")
    add_steps(doc, [
        (1, "Alterne o modo noturno", "Use o botão de tema no menu superior. A mudança é imediata e não interfere em relatórios ou dados."),
        (2, "Avalie se o reset é necessário", "O reset é uma ação de exceção, apropriada apenas para reinício controlado da operação."),
        (3, "Autentique-se como administrador", "O botão de reset aparece somente para o perfil autorizado."),
        (4, "Confirme a ação", "Leia a mensagem de confirmação e prossiga somente se os relatórios de backup necessários já estiverem salvos."),
        (5, "Importe uma nova base", "Após o reset, o histórico é reiniciado e a próxima importação passa a ser o primeiro ciclo."),
    ])
    add_note(doc, "Risco operacional", "Resetar importações é irreversível dentro da aplicação. Guarde o PDF Completo e a exportação Excel antes de qualquer limpeza planejada.", "FEF2F2")
    doc.add_page_break()

    add_title(doc, "13. Roteiro semanal recomendado", "Uma cadência para transformar informação em decisão.")
    add_steps(doc, [
        (1, "Receber e conferir", "Receba a planilha semanal e valide campos essenciais, integridade das datas e identificação de pedidos."),
        (2, "Importar", "Acesse com perfil administrativo e registre o novo arquivo semanal."),
        (3, "Ler o painel", "Observe índice de risco, itens vencidos, valor sob risco, alertas e pressão por filial."),
        (4, "Investigar exceções", "Filtre filiais e itens críticos; abra o histórico para confirmar a sequência de alterações."),
        (5, "Definir responsáveis", "Converta a fila de ação em responsáveis, prazos e encaminhamentos internos ou com fornecedores."),
        (6, "Reportar", "Emita PDF Executivo para liderança e Excel completo ou somente alterados para as áreas operacionais."),
        (7, "Preservar a sequência", "Não sobrescreva ciclos anteriores; a próxima planilha deve ser importada como um novo evento."),
    ])
    add_note(doc, "Disciplina de governança", "O valor do sistema aumenta quando o upload e a reunião de análise ocorrem em uma cadência estável. A repetição permite observar tendências, e não somente eventos isolados.")
    doc.add_page_break()

    add_title(doc, "14. Dúvidas frequentes e glossário", "Referência rápida para uso seguro e consistente.")
    add_table(doc, ["Dúvida", "Orientação"], [
        ("Por que não há alteração?", "Confirme se Ship To, Item e Customer PO correspondem ao mesmo pedido entre os dois uploads e se a previsão realmente mudou."),
        ("Por que o item saiu da base ativa?", "Se ele estava no upload anterior e não apareceu no atual, a regra o classifica como entregue."),
        ("O Excel traz todas as mudanças?", "Sim. A Base Operacional resume as datas e o Histórico de Alterações detalha cada transição exportada."),
        ("Como exportar apenas as exceções?", "Use Exportar alterados; o arquivo mantém filtros ativos e inclui itens com alteração de previsão."),
        ("Quando usar o PDF Completo?", "Quando for necessário levar todos os detalhes de alertas, tendências, base e mapa de alterações."),
    ], [5.4, 10.0])
    doc.add_heading("Glossário", level=2)
    add_table(doc, ["Termo", "Definição"], [
        ("Carteira ativa", "Itens presentes no upload mais recente e ainda em acompanhamento."),
        ("Instabilidade", "Recorrência de alterações na previsão de entrega de um item."),
        ("Valor sob risco", "Exposição financeira associada aos itens que atendem aos critérios de risco e alerta."),
        ("Previsão anterior", "Última data registrada antes da previsão atual do item."),
        ("Mapa de Alterações", "Visão que evidencia cada transição de previsão registrada no histórico."),
    ], [4.2, 11.2])
    doc.add_page_break()

    add_title(doc, "Referências e rastreabilidade do manual", "Fontes internas usadas para este documento.")
    add_body(doc, "As funcionalidades, fórmulas, controles e telas descritos neste manual foram verificados na aplicação Open Order Control, no código-fonte do projeto e no preview de referência validado em 15/08/2026. Não foram utilizados dados externos para caracterizar os indicadores do sistema.")
    add_table(doc, ["Referência", "Descrição"], [
        ("[1] Aplicação Open Order Control", "Dashboard, abas, filtros, login administrativo, relatórios e exportações validados no checkpoint 5f125413."),
        ("[2] Código do projeto", "Lógicas de comparação de uploads, risco executivo, alertas, normalização de filiais, itens entregues e exportação Excel."),
        ("[3] Capturas de tela", "Preview da aplicação em 15/08/2026; arquivos listados em executive_materials/assets/asset_sources.md."),
    ], [4.4, 11.0])
    add_note(doc, "Manutenção do manual", "Sempre que houver mudanças no fluxo de importação, nas métricas, nas permissões ou nos relatórios, atualize este manual e as capturas de tela correspondentes.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_manual()
