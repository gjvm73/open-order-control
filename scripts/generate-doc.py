import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette - Swiss Style (Charcoal & Red accent)
    COLOR_PRIMARY = RGBColor(26, 26, 26)     # #1A1A1A Charcoal
    COLOR_SECONDARY = RGBColor(180, 0, 0)   # #B40000 Swiss Red
    COLOR_TEXT = RGBColor(51, 51, 51)         # #333333 Dark Gray
    COLOR_MUTED = RGBColor(102, 102, 102)     # #666666 Muted Gray

    # Helper: set cell shading
    def set_cell_background(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Helper: set cell margins
    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Title Style
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_sub = title_p.add_run("MANUAL CONCEITUAL E METODOLOGIA EXECUTIVA\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    run_title = title_p.add_run("Open Order Control: Métricas, Indicadores e Tomada de Decisão")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(16)
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = COLOR_MUTED
    r_div.font.size = Pt(9)

    # Helper for Headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_paragraph(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT
        return p

    def add_callout(text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F5F5F5")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        # Left border red
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="B40000"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = COLOR_PRIMARY
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- Section 1 ---
    add_heading_1("1. Visão Geral e Concepção do Sistema")
    add_paragraph(
        "O Open Order Control é uma ferramenta web de inteligência logística e suprimentos projetada para o "
        "acompanhamento rigoroso de pedidos em aberto (open orders). O objetivo central do sistema é automatizar a comparação "
        "semanal de planilhas de pedidos enviadas por fornecedores ou ERPs, rastreando desvios nas previsões de entrega, "
        "quantificando alterações e fornecendo uma base sólida para a tomada de decisão executiva."
    )
    add_paragraph(
        "Em ambientes industriais e de cadeia de suprimentos, a oscilação constante nas datas de entrega compromete o planejamento "
        "de produção e o atendimento ao cliente. O sistema resolve essa dor ao consolidar o histórico de cada item com base em uma "
        "chave estável composta por Endereço (Ship To), Código do Item e Ordem de Compra do Cliente (Customer PO)."
    )

    # --- Section 2 ---
    add_heading_1("2. O Ciclo Semanal e a Lógica de Comparação")
    add_paragraph(
        "O fluxo operacional baseia-se no upload semanal de planilhas nos formatos Excel (.xlsx ou .xls). A cada nova importação, "
        "o sistema processa milhares de linhas em segundos, realizando um cruzamento inteligente com o histórico acumulado."
    )
    add_paragraph(
        "Para cada linha da nova planilha, o sistema localiza o item correspondente e compara a data de previsão informada ("
        "coluna 'Previsão') com a última previsão registrada no banco de dados. Caso haja divergência na data, a ocorrência "
        "é gravada na tabela de histórico, o contador de alterações do item é incrementado e a nova data passa a ser o padrão vigente."
    )
    add_callout(
        "Regra de Ouro do Histórico: O sistema diferencia rigorosamente o comportamento do ciclo mais recente (quantos itens "
        "mudaram no último upload) do acumulado histórico (quantas vezes o item sofreu alterações ao longo de toda a sua vida útil)."
    )

    # --- Section 3 ---
    add_heading_1("3. Dicionário de Indicadores e Fórmulas de Cálculo")
    add_paragraph(
        "Para apoiar a diretoria e os gestores de suprimentos, o painel gerencial consolida os dados em métricas acionáveis. "
        "Abaixo estão detalhadas as fórmulas matemáticas e lógicas que alimentam os cartões e gráficos do sistema:"
    )

    # Table of Metrics
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Métrica Executiva", "Fórmula / Lógica de Cálculo", "Interpretação para Decisão"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1A1A1A")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)

    metrics_data = [
        ("Estabilidade do Último Ciclo", "max(0, 100 - (Itens Alterados no Último Upload / Total de Itens Ativos) × 100)", "Mede o percentual da carteira que manteve os prazos inalterados na virada de semana mais recente."),
        ("Índice de Risco Executivo", "min(100, (Taxa Alteração × 0.35) + (Taxa Vencidos × 0.25) + (Taxa Sem Fornecedor × 0.20) + (Taxa Prioridade Alta × 0.10) + (Taxa Financeira Sob Risco × 0.10))", "Nota de 0 a 100 que pondera a pressão operacional da carteira ou filial selecionada."),
        ("Valor Sob Risco Acumulado", "Somatório do Preço Estendido (Extended Price) de todos os itens que possuem 1 ou mais alterações registradas.", "Volume financeiro alocado em pedidos que já sofreram instabilidade de cronograma."),
        ("Taxa de Risco da Carteira", "(Itens com Alterações Acumuladas / Total de Itens Ativos) × 100", "Percentual histórico de itens que apresentaram pelo menos uma oscilação de prazo desde o primeiro upload.")
    ]

    for row_idx, data in enumerate(metrics_data):
        row_cells = table.add_row().cells
        bg = "F9F9F9" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate(data):
            row_cells[i].text = text
            set_cell_background(row_cells[i], bg)
            set_cell_margins(row_cells[i], top=120, bottom=120, left=150, right=150)
            p = row_cells[i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- Section 4 ---
    add_heading_1("4. O Índice de Risco Executivo: Detalhadamente")
    add_paragraph(
        "O Índice de Risco Executivo é o coração do centro de comando do sistema. Ele resume a saúde da carteira em um único "
        "número ponderado, permitindo que o gestor saiba imediatamente se a operação está controlada ou sob forte estresse."
    )
    add_paragraph(
        "A composição do índice reflete cinco pilares logísticos com pesos calibrados para a realidade industrial:"
    )
    add_paragraph(
        "• Alterações Recentes (Peso 35%): Avalia a volatilidade imediata no último upload semanal.\n"
        "• Previsões Vencidas (Peso 25%): Mede o volume de itens cuja data de entrega já passou e ainda não foram faturados.\n"
        "• Itens Sem Fornecedor (Peso 20%): Identifica gargalos críticos onde o abastecimento está desamparado.\n"
        "• Prioridade Alta (Peso 10%): Pondera pedidos marcados como críticos pelo cliente ou pela operação.\n"
        "• Exposição Financeira (Peso 10%): Considera a proporção do valor financeiro afetado em relação ao total da carteira."
    )
    add_paragraph(
        "As faixas de classificação do índice orientam a postura gerencial: Abaixo de 25 pontos = Controlado (rotina de acompanhamento); "
        "Entre 25 e 49.9 pontos = Atenção (cobrança de fornecedores e revisão de prazos); 50 pontos ou mais = Crítico (reunião de "
        "alinhamento com diretoria e planos de contingência)."
    )

    # --- Section 5 ---
    add_heading_1("5. Alertas, Tendências e Consolidação por Filial")
    add_paragraph(
        "O sistema conta com um motor de alertas configurável pelo usuário. Definindo um limiar de dias (por exemplo, 7 dias), "
        "o sistema categoriza automaticamente as variações de prazo em duas severidades:"
    )
    add_paragraph(
        "1. Alerta de Atenção: Variações moderadas de prazo que exigem monitoramento preventivo.\n"
        "2. Alerta Crítico: Adiamentos severos que ultrapassam o limiar configurado, impactando diretamente o compromisso com o cliente."
    )
    add_paragraph(
        "Além disso, a consolidação por filial (Ship To) permite que empresas com múltiplas unidades ou centros de distribuição "
        "filtresm os dados para analisar a pressão específica de cada filial, garantindo que problemas locais não fiquem mascarados "
        "pela média geral da corporação."
    )

    # Save document
    file_path = "/home/ubuntu/open-order-control/Open_Order_Control_Manual_Conceitual.docx"
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    create_document()
